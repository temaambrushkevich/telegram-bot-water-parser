import time
import shutil
import random
import datetime

import urllib.request
from selenium import webdriver
from selenium.webdriver.common.by import By
from selenium.webdriver.support.ui import WebDriverWait
from selenium.webdriver.support import expected_conditions as EC
import chromedriver_binary

from openpyxl import load_workbook  # для записи в excel
from openpyxl.worksheet.dimensions import ColumnDimension, DimensionHolder
from openpyxl.utils import get_column_letter
from openpyxl.styles import PatternFill, Border, Side, Alignment, Font

from docx2pdf import convert
import docx
from docx.shared import Cm, Pt
import aspose.words as aw


# ПАРСИНГ ЦЕН
def parsing_price():
    FILE_NAME_PRICE = "цены за всё время.xlsx"

    # проверить excel файл на наличие листа, если есть - удалить
    def check_sheet_in_excel():
        file_name = FILE_NAME_PRICE
        wbook = load_workbook(file_name)

        now = datetime.datetime.now()
        date = now.strftime("%d.%m.%Y")
        sheet_name = date  # дата с компьютера
        if sheet_name in wbook:
            wbook.remove(wbook[sheet_name])
            wbook.save(filename=file_name)
        wbook.save(filename=file_name)
    # функция для добавления данных в Excel
    def add_excel(name, price_one_b, price_two_b, price_tara):
        colors = ["FFFFE4B5", "FFFFE4E1", "7FFFD4", "E0EEEE", "FFE4C4", "FFEBCD", "98F5FF", "FF9912", "66CD00",
                  "FF7F24",
                  "FFF8DC", "FF8C00", "BF3EFF", "F2F2F2", "ADFF2F", "FF6A6A", "FFF68F", "FFF0F5", "BFEFFF", "FFA07A",
                  "A4D3EE",
                  "E066FF", "B3EE3A", "BBFFFF", "FFBBFF", "FFC1C1"]

        file_name = FILE_NAME_PRICE
        wbook = load_workbook(file_name)

        # стиль шрифта для названий
        font_name = Font(size=12, bold=True)

        index = len(wbook.sheetnames)

        now = datetime.datetime.now()
        date = now.strftime("%d.%m.%Y")
        sheet_name = date  # дата с компьютера

        def function():
            # ширина колонок
            def len_columns():
                dim_holder = DimensionHolder(worksheet=ws)
                for col in range(ws.min_column, ws.max_column + 1):
                    # width=len(cell_value) * 1.23
                    dim_holder[get_column_letter(col)] = ColumnDimension(ws, min=col, max=col, width=25)
                ws.column_dimensions = dim_holder

            # узнать, какая колонка свободна
            def check_empty_cell():
                empty_cell_randc = []

                value = 1
                row = 1
                column = 1
                while (value != None):
                    value = (ws.cell(row=row, column=column)).value
                    if value == None:
                        break
                    else:
                        column += 1

                empty_cell_randc.append(row)
                empty_cell_randc.append(column)
                return empty_cell_randc

            # добавить значения в лист
            def add_values():
                call_name_r1 = ws.cell(row=call_row, column=call_col).column_letter + str(
                    ws.cell(row=call_row, column=call_col).row)
                ws.cell(row=call_row, column=call_col, value=name).font = font_name
                ws.cell(row=call_row, column=call_col, value=name).border = border_h
                ws.cell(row=call_row, column=call_col).alignment = Alignment(wrap_text=True, horizontal='center',
                                                                             vertical='center')  # включаем перенос строк и выравнивание для ячеек в первой строке

                # ws.cell(row=call_row + 1, column=call_col, value=price_one_b).number_format = '#,##0 ₽'
                ws.cell(row=call_row + 1, column=call_col, value=price_one_b).border = border_2

                ws.cell(row=call_row + 2, column=call_col, value=price_two_b).border = border_2

                # формула  пример: =ЕСЛИ(B3 = 0; 0;1-'02.18'!B3/B3)     АНГЛОЯЗЫЧНЫЙ ВАРИАНТ, там не ; а ,
                # =ЕСЛИ('31.10.2022'!B1='06.11.2022'!B1;ЕСЛИ(B3=0; 0; 1-'30.10.2022'!B3/B3);0)
                # название ячейки строка 3
                call_name_r3 = ws.cell(row=call_row + 2, column=call_col).column_letter + str(
                    ws.cell(row=call_row + 2, column=call_col).row)
                # print(call_name_r3)
                # print("=IF(" + '\'' + name_last_sheet + "\'!" + call_name_r1 + "=\'" + sheet_name + "\'!" + call_name_r1 + ', IF(' + call_name_r3 + "=0, 0, 1-"+"\'" + name_last_sheet + "\'" + "!" + call_name_r3 + "/" + call_name_r3 + ")" + ", 0)")

                ws.cell(row=call_row + 3, column=call_col,
                        value="=IF(" + '\'' + name_last_sheet + "\'!" + call_name_r1 + "=\'" + sheet_name + "\'!" + call_name_r1 + ', IF(' + call_name_r3 + "=0, 0, 1-" + "\'" + name_last_sheet + "\'" + "!" + call_name_r3 + "/" + call_name_r3 + ")" + ", 0)")
                ws.cell(row=call_row + 3, column=call_col).border = border_2
                ws.cell(row=call_row + 3, column=call_col).number_format = "0.0%"

                ws.cell(row=call_row + 4, column=call_col, value=price_tara).border = border_3

            # настроить высоту первой строки
            ws.row_dimensions[1].height = 64

            call_row = check_empty_cell()[0]
            call_col = check_empty_cell()[1]

            # вписываем текст в самые левые ячейки
            if call_row == 1 and call_col == 1:
                # назначаем цвет первой колонке
                h_fill = PatternFill(start_color='FFE4C4',
                                     end_color='FFE4C4',
                                     fill_type='solid')
                ws['A1'].fill = h_fill
                ws['A2'].fill = h_fill
                ws['A3'].fill = h_fill
                ws['A4'].fill = h_fill
                ws['A5'].fill = h_fill

                # вписываем текст в первые пять ячеек первой колонки
                ws.cell(row=call_row, column=call_col, value="Название").font = font_name
                ws.cell(row=call_row, column=call_col).border = border_h
                ws.cell(row=call_row, column=call_col).alignment = Alignment(wrap_text=True, horizontal='center',
                                                                             vertical='center')  # включаем перенос строк и выравнивание для ячеек в первой строке

                ws.cell(row=call_row + 1, column=call_col, value="19л 1шт").border = border_h
                ws.cell(row=call_row + 1, column=call_col).font = font_name

                ws.cell(row=call_row + 2, column=call_col, value="19л от 2х").border = border_h
                ws.cell(row=call_row + 2, column=call_col).font = font_name

                ws.cell(row=call_row + 3, column=call_col, value="Изм. цены").border = border_h
                ws.cell(row=call_row + 3, column=call_col).font = font_name

                ws.cell(row=call_row + 4, column=call_col, value="Тара").border = border_h
                ws.cell(row=call_row + 4, column=call_col).font = font_name

                call_col += 1

            add_values()

            len_columns()
            ws.column_dimensions['A'].width = 15
            ws.alignment = Alignment(wrap_text=True)
            wbook.save(file_name)

        def color_sheet():
            # цвет листа внизу
            sheet_properties = ws.sheet_properties
            random_num = random.randint(0, len(colors) - 1)
            sheet_properties.tabColor = colors[random_num]

        # настроить стили граней ячеек
        bd = Side(border_style='medium')
        bd2 = Side(border_style='thin')
        border_h = Border(left=bd, top=bd, right=bd, bottom=bd)
        border_2 = Border(left=bd, top=bd2, right=bd, bottom=bd2)
        border_3 = Border(left=bd, top=bd2, right=bd, bottom=bd)

        if sheet_name in wbook:
            print("Такая страница есть")
            ws = wbook[sheet_name]  # лист
            # узнаём имя листа с предыдущими измерениями
            name_last_sheet = wbook.sheetnames[index - 2]
            color_sheet()
            function()

        else:
            print("Такой страницы нет")
            wbook.create_sheet(index=index, title=sheet_name)
            ws = wbook[sheet_name]  # лист
            name_last_sheet = wbook.sheetnames[index - 1]
            color_sheet()
            function()
    # добавить цену на товар в сводный лист в Excel
    def add_listsummary(tovar_name, price):
        file_name = FILE_NAME_PRICE
        sheet_name = "Общая от 2-х"

        wbook = load_workbook(file_name)
        sheet = wbook[sheet_name]
        print("OK, лист найден")

        now = datetime.datetime.now()
        date = now.strftime("%d.%m.%Y")
        print("NOW DATE = " + date)

        # проверка есть ли колонка с названием товара:  return(True / False)
        def check_item_name():
            check_item_name = True
            print("Ищем колонку с названием товара...")

            value = "name"
            row = 1
            column = 2
            while (value != tovar_name and value != None):
                value = (sheet.cell(row=row, column=column)).value
                if value == tovar_name:
                    return check_item_name
                else:
                    column += 1

            check_item_name = False
            return check_item_name

        # ПОИСК НУЖНОЙ КОЛОНКИ
        def find_desired_column():
            print("Ищем нужную колонку")

            value = "name"
            row = 1
            column = 2
            while (value != tovar_name and value != None):
                value = (sheet.cell(row=row, column=column)).value
                if value == tovar_name:
                    print("Нужная колонка найдена, её имя: ", value, '\n')
                    return column
                else:
                    column += 1

            return column

        # ПОИСК НУЖНОЙ СТРОЧКИ В СТОЛБЦЕ
        def find_desired_row():
            print("Ищем нужную строчку...")
            row = 2
            column = 1
            value = (sheet.cell(row=row, column=column)).value

            while (value != date and value != None):
                value = (sheet.cell(row=row, column=column)).value
                #print(value, '\n')
                if (str(value) == date):
                    return row
                # если строчка с сегодняшней датой не найдена - записать ее в первый столбец
                if (value == None):
                    sheet.cell(row=row, column=1, value=date)
                    return row
                row += 1

            sheet.cell(row=row, column=1, value=date)
            return row

        item_column = find_desired_column()
        empty_row = find_desired_row()

        # если колонка с таким именем товара уже есть
        if check_item_name():
            print("Такой товар есть : )")
            # Добавляем значение в ячейку нужного столбца
            sheet.cell(row=empty_row, column=item_column, value=price)
        # если такой колонки нет
        else:
            print("Такого товара не было")
            # Добавляем значение в ячейку нужного столбца
            sheet.cell(row=1, column=item_column - 1, value=tovar_name)
            # добавляем название товара в первую строчку таблицы
            sheet.cell(row=empty_row, column=item_column - 1, value=price)

        wbook.save(file_name)

    def close_browser():
        time.sleep(1)
        browser.quit()

    def parsing_lider():
        def get_page():
            # chromedriver должен быть ТОЙ ЖЕ ВЕРСИИ ЧТО И установленный CHROME
            browser.get("https://artvod.ru/product-category/oborudovanie/")
            print("===================== Парсится сайт artvod.ru =====================")
        def pars_water():
            # ПАРСИМ ЦЕНЫ
            def pars_tara():
                # парсим название позиций во вкладке оборудование
                item_class_name = "vitrina_name"
                WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, item_class_name)))
                items = browser.find_elements(By.CLASS_NAME, item_class_name)

                n = 0
                number_tara = 0
                array_items_names = []

                for item in items:
                    array_items_names.append(item.text)
                    if item.text.find("Тара ПЭТ") != -1:
                        number_tara = n
                    n += 1

                # парсим цены позиций
                price_class_name = "vitrina_price_fon"
                items_prices = browser.find_elements(By.CLASS_NAME, price_class_name)
                array_prices = []
                for item in items_prices:
                    array_prices.append(item.text)

                price_tara = int((array_prices[number_tara]).replace('р.', ''))
                print("Стоимость тары: ", price_tara)  # выводим стоимость ТАРЫ ПЭТ

                return price_tara
            def get_item_page(n):
                browser.get("https://artvod.ru/product-category/pityevaya-voda/")
                items_class_name = "vitrina_name"
                WebDriverWait(browser, 200).until(EC.presence_of_element_located((By.CLASS_NAME, items_class_name)))
                time.sleep(2)
                items = browser.find_elements(By.CLASS_NAME, items_class_name)
                items[n].click()
                parsing_price()
            def parsing_price():
                values_list = []

                WebDriverWait(browser, 200).until(EC.presence_of_element_located((By.CLASS_NAME, "tovar_name")))
                tovar_name = (browser.find_element(By.CLASS_NAME, "tovar_name")).text
                print(tovar_name)

                values_list.append(tovar_name)  # далее в excel

                WebDriverWait(browser, 200).until(EC.presence_of_element_located((By.CLASS_NAME, "update-price")))
                el_price_one_bottle = browser.find_element(By.CLASS_NAME, "update-price")
                price_one_bottle = int(el_price_one_bottle.text)
                print("Цена при покупке одной бутылки: ", price_one_bottle)

                values_list.append(price_one_bottle)  # далее в excel

                plus_elem = browser.find_element(By.CLASS_NAME, "my_plus")
                plus_elem.click()
                time.sleep(4)

                el_price_two_bottle = browser.find_element(By.CLASS_NAME, "update-price")
                price_two_bottle = int(el_price_two_bottle.text) / 2
                print("Цена за штуку при покупке двух бутылок: ", price_two_bottle, end="\n\n")

                values_list.append(price_two_bottle)  # далее в excel
                values_list.append(price_tara)  # далее в excel

                add_listsummary(values_list[0], values_list[2])
                add_excel(values_list[0], values_list[1], values_list[2], values_list[3])

            price_tara = pars_tara()
            get_item_page(0)
            get_item_page(1)
            get_item_page(2)

        print("ПАРСИМ САЙТ ЛИДЕР")
        get_page()
        pars_water()
    def parsing_niagara():
        def get_page():
            browser.get("https://niagara74.ru/")
            print("===================== Парсится сайт niagara74.ru =====================")
            time.sleep(2)
        def parsing():
            values_list = []

            def find_name():
                class_name = "name"
                WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, class_name)))
                water_name = browser.find_element(By.CLASS_NAME, class_name)
                print(water_name.text)

                values_list.append(water_name.text)  # далее в excel
            def find_price():
                price_class_name = "price-col"
                WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, price_class_name)))
                price_item = browser.find_element(By.CLASS_NAME, price_class_name)
                print("[find price OK]")
                price_text = (price_item.text).split("\n")

                print(price_text[0], '\n')
                print(price_text[1], '\n')

                price_one_b = int(price_text[0].replace('₽', ""))
                price_two_b = int(price_text[1].replace('₽', ""))
                print("Цена при покупке одной бутылки: ", price_one_b)
                print("Цена за штуку при покупке двух бутылок: ", price_two_b)

                values_list.append(price_one_b)  # далее в excel
                values_list.append(price_two_b)  # далее в excel

                # далее спарсим цену тары
                price_class_name = "price"
                items_prices = browser.find_elements(By.CLASS_NAME, price_class_name)
                price_tara = int((items_prices[0].text).replace('₽', ""))
                print("Стоимость тары: ", price_tara)

                values_list.append(price_tara)  # далее в excel
                add_excel(values_list[0], values_list[1], values_list[2], values_list[3])
                add_listsummary(values_list[0], values_list[2])

            find_name()
            find_price()

        print("ПАРСИМ САЙТ НИАГАРА")
        get_page()
        parsing()
    def parsing_voda174_krystal():
        def get_page():
            browser.get("https://voda174.ru/")
            print("===================== Парсится сайт voda174.ru =====================")
        def parsing():
            # при покупке одной бутылки
            price_xpath1 = "/html/body/div[1]/div[2]/div[4]/div[3]/div/div[7]/div/div/a"
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.XPATH, price_xpath1)))
            print("[find OK]")
            price_elem1 = browser.find_element(By.XPATH, price_xpath1)
            price_one_b = int(((price_elem1.text).replace("руб", "")).replace("=", ""))
            print("Цена при покупке одной бутылки: ", price_one_b)

            # при покупке двух бутылок
            price_xpath2 = "/html/body/div[1]/div[2]/div[4]/div[3]/div/div[20]/div/div/a"
            price_elem2 = browser.find_element(By.XPATH, price_xpath2)
            price_two_b = int(((price_elem2.text).replace("руб", "")).replace("=", "")) / 2
            print("Цена за штуку при покупке двух бутылок: ", price_two_b)

            # тара
            price_xpath3 = "/html/body/div[1]/div[2]/div[7]/div[1]/div/div[2]/div/div/div/div[1]/div[1]/div/div[3]"
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.XPATH, price_xpath3)))
            print("price на тару найден")
            price_elem3 = browser.find_element(By.XPATH, price_xpath3)
            price_tara = int((price_elem3.text).replace("руб.", ""))
            print("Стоимость тары: ", price_tara)

            # в excel
            add_excel("voda174.ru кристальная", price_one_b, price_two_b, price_tara)
            add_listsummary("voda174.ru кристальная", price_two_b)

        print("ПАРСИМ САЙТ КРИСТАЛЬНАЯ(voda174.ru)")
        get_page()
        parsing()
    def parsing_living_drop():
        def get_page():
            browser.get("https://xn--80aaepkoi5a5le.xn--p1ai/catalog/voda/")
            print("===================== Парсится сайт живаякапля.рф =====================")
        def close_reklama():
            # close reklam-box
            WebDriverWait(browser, 10).until(EC.presence_of_element_located((By.XPATH, "/html/body/div[20]/div[2]/a")))
            print("==Реклама на живаякапля есть==")
            exit_reklama = browser.find_element(By.XPATH, "/html/body/div[20]/div[2]/a")
            exit_reklama.click()
            print("==Реклама закрыта==")
        def scrolling_end():
            # ПРОКРУЧИВАНИЕ СТРАНИЦЫ ВНИЗ ДО КОНЦА ОТЗЫВОВ
            SCROLL_PAUSE_TIME = 2
            # Получаем высоту прокрутки
            last_height = browser.execute_script("return document.body.scrollHeight")
            while True:
                # Прокручивание вниз
                browser.execute_script("window.scrollTo(0, document.body.scrollHeight);")

                # Ждем загрузки страницы
                time.sleep(SCROLL_PAUSE_TIME)

                # Рассчитываем новую высоту прокрутки и сравниваем с последней высотой прокрутки
                new_height = browser.execute_script("return document.body.scrollHeight")
                if new_height == last_height:
                    break
                last_height = new_height
        def parsing():
            scrolling_end()

            # переход на позицию с водой
            path1 = "/html/body/div[11]/div[7]/div[2]/div/div/div/div/div[2]/div[1]/div/div/div/div[2]/div[3]/div[1]/div[2]/div[1]/div[2]/div/div[2]/a/span/span"
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.XPATH, path1)))
            print("==OK, позиция найдена==")
            el_bottle = browser.find_element(By.XPATH, path1)
            el_bottle.click()
            print("==OK, клик произведен==")

            # цена при покупке 1 бутылки без тары
            path_price = "/html/body/div[11]/div[8]/div[2]/div/div/div/div/div/div/div/div[2]/div/div[3]/div[2]/div[1]/div/span/span[1]"
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.XPATH, path_price)))
            print("==OK, цена найдена==")
            el_price_one_b_nt = browser.find_element(By.XPATH, path_price)
            price_one_b_nt = el_price_one_b_nt.text

            # выбираем: одна бутылка и есть тара
            path2 = "/html/body/div[11]/div[8]/div[2]/div/div/div/div/div/div/div/div[2]/div/div[6]/div/div[1]/div/div/div/div/div/div[1]/div/ul/li[2]/span"
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.XPATH, path2)))
            print("==Кнопка найдена==")
            el_button = browser.find_element(By.XPATH, path2)
            el_button.click()
            print("==Клик ОК==")

            time.sleep(4)

            # цена при покупке 1 бутылки
            path_price1 = "/html/body/div[11]/div[8]/div[2]/div/div/div/div/div/div/div/div[2]/div/div[3]/div[1]/div[1]/div[1]/span[1]"
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.XPATH, path_price1)))
            print("==OK, цена найдена==")
            el_price_one_b = browser.find_element(By.XPATH, path_price1)
            price_one_b = el_price_one_b.text

            # Цена тары
            price_tara = int(price_one_b_nt) - int(price_one_b)


            print("Цена при покупке, если нет тары: " + str(price_one_b_nt))
            print("Цена при покупке от 1 до 3 бутылок и если есть тара: " + str(price_one_b))
            print("Цена тары: " + str(price_tara))
            # в excel
            add_excel("Живая капля", int(price_one_b), int(price_one_b), int(price_tara))
            add_listsummary("Живая капля", int(price_one_b))


        print("ПАРСИМ ЖИВАЯ КАПЛЯ")
        get_page()
        try:
            close_reklama()
        except:
            print("==Рекламы нет==")
        parsing()
    def parsing_oazis74():
        def get_page():
            browser.get("https://www.74mv.ru/katalog/gornyj-oazis")
            print("===================== Парсится сайт www.74mv.ru =====================")

        def parsing():
            # парсим цену за воду
            WebDriverWait(browser, 20).until(
                EC.presence_of_element_located((By.LINK_TEXT, "Вода питьевая \"Горный Оазис\" негазированная 19,0л")))
            print("ok")
            name = browser.find_element(By.LINK_TEXT, "Вода питьевая \"Горный Оазис\" негазированная 19,0л")
            print(name.text)
            el_price = browser.find_element(By.CLASS_NAME, "PriceunitPrice")
            price_two_b = int(((el_price.text).replace("Цена / шт:", "")).replace(",00 руб.", ""))
            print("Цена за штуку при покупке двух бутылок: ", price_two_b)  # МЕНЬШЕ ОДНОЙ НЕЛЬЗЯ

            # парсим цену за тару
            text_element = browser.find_element(By.CLASS_NAME, "product_s_desc")
            index_price = (text_element.text).find("Залоговая стоимость тары")
            price_tara = int(
                (text_element.text)[index_price + 25] + (text_element.text)[index_price + 26] + (text_element.text)[
                    index_price + 27])
            print("Стоимость тары: ", price_tara)

            # excel
            add_excel("Горный Оазис", price_two_b, price_two_b, price_tara)
            add_listsummary("Горный Оазис", price_two_b)

        print("ПАРСИМ ОАЗИС")
        get_page()
        parsing()
    def parsing_vlasovkluch():
        def get_page():
            browser.get("http://vlasovkluch.ru/cat/")
            print("===================== Парсится сайт vlasovkluch.ru =====================")

        def parsing():
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "txt-item-card")))
            products = browser.find_element(By.CLASS_NAME, "txt-item-card")
            products.click()

            water_bot_cname = "price-item-product"
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, water_bot_cname)))
            el_water_bottles_prices = browser.find_elements(By.CLASS_NAME, water_bot_cname)
            arr_prices = []
            j = 0
            for item in el_water_bottles_prices:
                price = int((item.text).replace("Р", ""))
                if j <= 1:
                    arr_prices.append(price)
                j += 1

            water_name_cname = "title-item-product"
            el_water_bottles_names = browser.find_elements(By.CLASS_NAME, water_name_cname)
            arr_names = []
            n = 0
            for item in el_water_bottles_names:
                if n <= 1:
                    arr_names.append(item.text)
                n += 1

            print(arr_names)
            print("Цена при покупке одной бутылки: ")
            print(arr_prices)

            # excel
            add_excel(arr_names[0], arr_prices[0], arr_prices[0], 0)
            add_listsummary(arr_names[0], arr_prices[0])
            add_excel(arr_names[1], arr_prices[1], arr_prices[1], 0)
            add_listsummary(arr_names[1], arr_prices[1])

        print("ПАРСИМ ВЛАСОВ КЛЮЧ")
        get_page()
        parsing()
    def parsing_chebistok():
        def get_page():
            browser.get("https://chebistok.ru/")
            print("===================== Парсится сайт chebistok.ru =====================")
        def parsing():
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "price")))
            el_price = browser.find_element(By.CLASS_NAME, "price")
            price_two_bottles = int((el_price.text).replace("р./шт.", ""))
            print("Цена за штуку при покупке двух бутылок: ", price_two_bottles)

            add_excel("Чебаркульский исток", price_two_bottles, price_two_bottles, 0)
            add_listsummary("Чебаркульский исток", price_two_bottles)

        print("ПАРСИМ ЧЕБАРКУЛЬСКИЙ ИСТОК")
        get_page()
        parsing()
    def parsing_l_w():
        def get_page():
            browser.get("https://l-w.ru/catalog/voda/")
            print("===================== Парсится сайт l-w.ru =====================")
        def parsing():
            time.sleep(2)
            # закрываем рекламу
            try:
                WebDriverWait(browser, 200).until(EC.presence_of_element_located((By.CLASS_NAME, "fancybox-close-small")))
                close_reklama = browser.find_element(By.CLASS_NAME, "fancybox-close-small")
                close_reklama.click()
            except:
                print("----рекламного банера нет----")

            time.sleep(5)
            # парсим названия
            names_items_cname = "product-slide__title"
            WebDriverWait(browser, 200).until(EC.presence_of_element_located((By.CLASS_NAME, names_items_cname)))
            names_elements = browser.find_elements(By.CLASS_NAME, names_items_cname)
            print(names_elements[0].text)
            print(names_elements[1].text)

            # парсим цены
            el_prices = browser.find_elements(By.CLASS_NAME, "price")
            arr_price_two_bottles = []

            arr_price_two_bottles.append((int((el_prices[0].text).replace("₽", ""))))
            arr_price_two_bottles.append((int((el_prices[1].text).replace("₽", ""))))
            print("Цена за штуку при покупке двух бутылок: ")
            print(arr_price_two_bottles)

            # убавляем на 1 кол-во бутылок
            # 1-ый товар
            WebDriverWait(browser, 200).until(EC.presence_of_element_located((By.XPATH, "/html/body/div[1]/main/div/div/div[2]/div[1]/div[4]/div/button[1]")))
            elements_minus1 = browser.find_element(By.XPATH, "/html/body/div[1]/main/div/div/div[2]/div[1]/div[4]/div/button[1]")
            elements_minus1.click()
            time.sleep(2)
            # 2-ой товар
            WebDriverWait(browser, 200).until(EC.presence_of_element_located((By.XPATH, "/html/body/div[1]/main/div/div/div[2]/div[2]/div[4]/div/button[1]")))
            elements_minus2 = browser.find_element(By.XPATH, "/html/body/div[1]/main/div/div/div[2]/div[2]/div[4]/div/button[1]")
            elements_minus2.click()
            time.sleep(5)
            # парсим цены за 1 бутылку
            arr_price_one_bottles = []
            WebDriverWait(browser, 200).until(EC.presence_of_element_located((By.CLASS_NAME, "price")))
            el_prices2 = browser.find_elements(By.CLASS_NAME, "price")
            arr_price_one_bottles.append(int((el_prices2[0].text).replace("₽", "")))
            arr_price_one_bottles.append(int((el_prices2[1].text).replace("₽", "")))
            print("Цена при покупке одной бутылки: ")
            print(arr_price_one_bottles)

            add_excel("\"Люкс Вода\"" + names_elements[0].text, arr_price_one_bottles[0], arr_price_two_bottles[0], 0)
            add_excel("\"Люкс Вода\"" + names_elements[1].text, arr_price_one_bottles[1], arr_price_two_bottles[1], 0)
            add_listsummary("\"Люкс Вода\"" + names_elements[0].text, arr_price_two_bottles[0])
            add_listsummary("\"Люкс Вода\"" + names_elements[1].text, arr_price_two_bottles[1])

        print("ПАРСИМ ЛЮКС ВОДА")
        get_page()
        parsing()
    def parsing_loveplus():
        def get_page():
            browser.get("https://vodalubima.ru/")
            print("===================== Парсится сайт vodalubima.ru =====================")
        def scrolling_end():
            # ПРОКРУЧИВАНИЕ СТРАНИЦЫ ВНИЗ ДО КОНЦА ОТЗЫВОВ
            SCROLL_PAUSE_TIME = 2
            # Получаем высоту прокрутки
            last_height = browser.execute_script("return document.body.scrollHeight")
            while True:
                # Прокручивание вниз
                browser.execute_script("window.scrollTo(0, document.body.scrollHeight);")

                # Ждем загрузки страницы
                time.sleep(SCROLL_PAUSE_TIME)

                # Рассчитываем новую высоту прокрутки и сравниваем с последней высотой прокрутки
                new_height = browser.execute_script("return document.body.scrollHeight")
                if new_height == last_height:
                    break
                last_height = new_height
        def parsing():
            scrolling_end()
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, 't776__price-wrapper')))
            print("OK")
            price_elems = browser.find_elements(By.CLASS_NAME, "t776__price-wrapper")
            price_two_b = int((price_elems[0].text).replace("руб.", ""))

            price_tara = int((price_elems[2].text).replace("руб.", ""))
            print(price_two_b)
            print(price_tara)

            add_excel("\"Вода-любимая+\"", price_two_b, price_two_b, price_tara)
            add_listsummary("\"Вода-любимая+\"", price_two_b)

        print("ПАРСИМ ВОДА-ЛЮБИМАЯ+")
        get_page()
        parsing()
    def parsing_arx74_artenza():
        def get_page():
            browser.get("https://xn--74-6kc2a7b4a5b.xn--p1ai/page/19litrov")
            print("===================== Парсится сайт архыз74.рф =====================")
        def parsing():
            # ПАРСИМ ЦЕНЫ
            # АРХЫЗ-ВИТА
            xpath_vita1 = "/html/body/section[2]/div/table[1]/tbody/tr[3]/td[1]/strong"
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.XPATH, xpath_vita1)))
            print("OK")
            one_bottle_arx_vita = browser.find_element(By.XPATH, xpath_vita1)
            price_oneb_vita = int((one_bottle_arx_vita.text).replace("руб.", ""))
            print("Архыз-вита за 1 бутылку: ", price_oneb_vita)

            xpath_vita2 = "/html/body/section[2]/div/table[1]/tbody/tr[3]/td[2]/strong"
            two_bottle_arx_vita = browser.find_element(By.XPATH, xpath_vita2)
            price_twob_vita = int((two_bottle_arx_vita.text).replace("руб.", ""))
            print("Архыз-вита за 1 бутылку при покупке от 2-х: ", price_twob_vita)

            # АРТЕНЗА
            xpath_artenza1 = "/html/body/section[2]/div/table[3]/tbody/tr[3]/td[1]/strong"
            xpath_artenza2 = "/html/body/section[2]/div/table[3]/tbody/tr[3]/td[2]/strong[1]"
            one_bottle_artenza = browser.find_element(By.XPATH, xpath_artenza1)
            two_bottle_artenza = browser.find_element(By.XPATH, xpath_artenza2)

            price_oneb_artenza = int((one_bottle_artenza.text).replace("руб.", ""))
            price_twob_artenza = int((two_bottle_artenza.text).replace("руб.", ""))

            print("Артенза за 1 бутылку: ", price_oneb_artenza)
            print("Артенза за 1 бутылку при покупке от 2-х: ", price_twob_artenza)

            # ПАРСИМ ЦЕНУ ТАРЫ
            xpath_tara = "/html/body/section[2]/div/p[7]"
            tara_el = browser.find_element(By.XPATH, xpath_tara)
            price_tara = int(
                ((tara_el.text).replace("Упаковка 19 л. Залоговая стоимость тары - ", "")).replace("рублей.", ""))
            print("Стоимость тары: ", price_tara)

            add_excel("\"архыз-74\" Архыз-вита", price_oneb_vita, price_twob_vita, price_tara)
            add_excel("\"архыз-74\" Артенза", price_oneb_artenza, price_twob_artenza, price_tara)
            add_listsummary("\"архыз-74\" Архыз-вита", price_twob_vita)
            add_listsummary("\"архыз-74\" Артенза", price_twob_artenza)

        print("ПАРСИМ АРХЫЗ74, архыз, артенза")
        get_page()
        parsing()
    def parsing_aqua_mobil():
        def get_page():
            browser.get("https://aqua-mobil.ru/voda-9-19-litrov/")
            print("===================== Парсится сайт aqua-mobil.ru =====================")
        def parsing():
            xpath_checkbox = "/html/body/div[1]/div[1]/div[3]/div/div[3]/div[1]/div[1]/form/div[2]/div[2]/div[2]/div/div[2]/label/span"
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.XPATH, xpath_checkbox)))
            print("OK")
            checkbox = browser.find_element(By.XPATH, xpath_checkbox)
            checkbox.click()
            time.sleep(1)

            arr_prices_two = []
            arr_names = []

            # ПАРСИМ ЦЕНЫ ОТ 2-Х БУТЫЛОК
            print("ЦЕНЫ ОТ 2-Х БУТЫЛОК:")
            price_elements = browser.find_elements(By.CLASS_NAME, "price")
            n = 0
            for item in price_elements:
                if n != 0:
                    price = int((item.text).replace(" руб.", ""))
                    arr_prices_two.append(price)
                n += 1

            names_elements = browser.find_elements(By.CLASS_NAME, "top")
            for item in names_elements:
                arr_names.append((item.text).replace("\n", " "))

            for i in range(0, len(arr_prices_two)):
                print(arr_names[i], ":", arr_prices_two[i])

            # ПАРСИМ ЦЕНЫ ЗА 1 БУТЫЛКУ
            print("ЦЕНЫ ЗА ОДНУ БУТЫЛКУ:")
            elements_minus = browser.find_elements(By.CLASS_NAME, "good__minus")
            j = 1
            for item in elements_minus:
                if j != 3:
                    item.click()
                    time.sleep(2)
                    close_el = browser.find_element(By.XPATH, "/html/body/div[1]/div[7]/div/div[1]")
                    close_el.click()
                    time.sleep(4)
                j += 1

            arr_prices_one = []
            price_elements = browser.find_elements(By.CLASS_NAME, "price")
            n = 0
            for item in price_elements:
                if n != 0:
                    price = int((item.text).replace(" руб.", ""))
                    arr_prices_one.append(price)
                n += 1

            for i in range(0, len(arr_prices_one)):
                print(arr_names[i], ":", arr_prices_one[i])

            # добавляем в excel
            for i in range(len(arr_prices_one)):
                add_excel(arr_names[i], arr_prices_one[i], arr_prices_two[i], 0)
                add_listsummary(arr_names[i], arr_prices_two[i])

        print("ПАРСИМ АКВАМОБИЛЬ")
        get_page()
        parsing()

    check_sheet_in_excel()

    # для локального запуска хром
    options = webdriver.ChromeOptions()
    browser = webdriver.Chrome(chrome_options=options)      # в скрытом режиме
    browser.maximize_window()

    total_error = 0
    try:
        parsing_lider()
        print("\n")
    except:
        total_error+=1
        print("-------------Не удалось спарсить artvod.ru-------------")

    try:
        parsing_niagara()
        print("\n")
    except:
        total_error += 1
        print("-------------Не удалось спарсить niagara74.ru-------------")

    try:
        parsing_voda174_krystal()
        print("\n")
    except:
        total_error += 1
        print("-------------Не удалось спарсить voda174.ru-------------")

    try:
        parsing_living_drop()
        print("\n")
    except:
        total_error += 1
        print("-------------Не удалось спарсить живаякапля.рф-------------")

    try:
        parsing_oazis74()
        print("\n")
    except:
        total_error += 1
        print("-------------Не удалось спарсить www.74mv.ru-------------")

    try:
        parsing_vlasovkluch()
        print("\n")
    except:
        total_error += 1
        print("-------------Не удалось спарсить vlasovkluch.ru-------------")

    try:
        parsing_chebistok()
        print("\n")
    except:
        total_error += 1
        print("-------------Не удалось спарсить chebistok.ru-------------")

    try:
        parsing_l_w()
        print("\n")
    except:
        total_error += 1
        print("-------------Не удалось спарсить l-w.ru-------------")

    try:
        parsing_loveplus()
        print("\n")
    except:
        total_error += 1
        print("-------------Не удалось спарсить vodalubima.ru-------------")

    try:
        parsing_arx74_artenza()
        print("\n")
    except:
        total_error += 1
        print("-------------Не удалось спарсить архыз74.рф-------------")

    try:
        parsing_aqua_mobil()
        print("\n")
    except:
        total_error += 1
        print("-------------Не удалось спарсить aqua-mobil.ru-------------")

    close_browser()

    # Проверка на ошибки парсинга
    if total_error != 0:
        msg = " ================= Всего ошибок парсинга: " + str(total_error) + '\n' + "Попробуйте запустить парсер еще раз ================= "
        print(msg)
    else:
        print("Цены успешно спарсены")

# ПАРСИНГ АКЦИЙ
def parsing_stock():

    STOCK_FILE_NAME = "акции.docx"

    def close_browser():
        browser.close()

    # output text:
    def niagara():
        def get_page():
            browser.get("https://niagara74.ru/stock/")
            print("===================== Парсится сайт niagara74.ru =====================")

        def parsing():
            print("НИАГАРА")
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "object-fit")))
            time.sleep(5)
            items = browser.find_elements(By.CLASS_NAME, "object-fit")

            for item in items:
                stock_text = (item.find_element(By.CSS_SELECTOR, 'img')).get_attribute('alt')
                print(stock_text)
                add_txtstock_in_docx(stock_text)

        get_page()
        parsing()
    def Suite_water():
        def get_page():
            browser.get("https://l-w.ru/promo/")
            time.sleep(2)
            print("===================== Парсится сайт l-w.ru =====================")

        def parsing():
            try:
                WebDriverWait(browser, 20).until(
                    EC.presence_of_element_located((By.CLASS_NAME, "fancybox-close-small")))
                print("OK")
                close_button = browser.find_element(By.CLASS_NAME, "fancybox-close-small")
                close_button.click()
                print("OK")
            except:
                print("Рекламы нет")
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "stocks__item")))
            print("Люкс вода")
            images1 = browser.find_elements(By.CLASS_NAME, "stocks__item")
            for item in images1:
                print(item.text)
                add_txtstock_in_docx(str(item.text))
                time.sleep(0.5)
                print("\n")
            WebDriverWait(browser, 20).until(
                EC.presence_of_element_located((By.XPATH, "/html/body/div[1]/main/div/div/ul/li[4]/a")))
            list_button = browser.find_element(By.XPATH, "/html/body/div[1]/main/div/div/ul/li[4]/a")
            list_button.click()
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "stocks__item")))
            images2 = browser.find_elements(By.CLASS_NAME, "stocks__item")
            for item in images2:
                print(item.text)
                add_txtstock_in_docx(str(item.text))
                time.sleep(0.5)
                print("\n")

        get_page()
        parsing()
    def parsing_living_capla():
        def get_page():
            browser.get("https://xn--80aaepkoi5a5le.xn--p1ai/sale/")
            print("===================== Парсится сайт живаякапля.рф =====================")

        def parsing():
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "inner-text")))
            print("ЖИВАЯ КАПЛЯ")
            images = browser.find_elements(By.CLASS_NAME, "inner-text")

            for item in images:
                print(item.text)

                add_txtstock_in_docx(item.text)

                print("\n")

        get_page()
        parsing()
    def parsing_aqua_mobil():
        def get_page():
            browser.get("https://aqua-mobil.ru/")
            print("===================== Парсится сайт aqua-mobil.ru =====================")


            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.XPATH, "/html/body/div[1]/div[1]/div[1]/ul/li[8]/a")))
            sale_button = browser.find_element(By.XPATH, "/html/body/div[1]/div[1]/div[1]/ul/li[8]/a")
            sale_button.click()
        def parsing():
            print("Аква мобиль")
            i = 0
            while i != 3:
                WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "zakazblock")))
                sales_button = browser.find_elements(By.CLASS_NAME, "zakazblock")
                sales_button[i].click()
                WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "maxcontent")))
                sale = browser.find_element(By.CLASS_NAME, "maxcontent")
                print(sale.text)

                add_txtstock_in_docx(sale.text)

                WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "active")))
                sale_button = browser.find_element(By.CLASS_NAME, "active")
                sale_button.click()
                i = i + 1

        get_page()
        parsing()
    def artvod():
        def get_page():
            browser.get("https://artvod.ru/product-category/akzii/")
            print("===================== Парсится сайт artvod.ru =====================")
        def parsing():
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "vitrina_buy")))
            print("OK")

            buy_buttons = browser.find_elements(By.CLASS_NAME, "vitrina_buy")

            maxx = len(buy_buttons)

            for i in range(0, maxx):
                buy_buttons = browser.find_elements(By.CLASS_NAME, "vitrina_buy")
                buy_buttons[i].click()
                time.sleep(1)
                tovar_opisanie = browser.find_element(By.CLASS_NAME, "tovar_opisanie")
                print(tovar_opisanie.text)

                add_txtstock_in_docx(tovar_opisanie.text)

                browser.get("https://artvod.ru/product-category/akzii/")
                time.sleep(1)

        get_page()
        parsing()
    def crystal_water():
        def get_page():
            browser.get("https://voda174.ru/")
            print("===================== Парсится сайт voda174.ru =====================")

        def parsing():
            print("Кристальная вода")
            WebDriverWait(browser, 20).until(
                EC.presence_of_element_located((By.XPATH, "/html/body/div[1]/div[2]/div[10]/div[3]/div/div[1]")))
            sale = browser.find_element(By.XPATH, "/html/body/div[1]/div[2]/div[10]/div[3]/div/div[1]")
            print(sale.text)

            add_txtstock_in_docx(sale.text)

        get_page()
        parsing()
    def mountain_oasis():
        def get_page():
            browser.get("https://www.74mv.ru/")
            WebDriverWait(browser, 20).until(
                EC.presence_of_element_located((By.XPATH, "/html/body/div[1]/header/nav/div[1]/div/ul/li[4]")))
            sale_button = browser.find_element(By.XPATH, "/html/body/div[1]/header/nav/div[1]/div/ul/li[4]")
            sale_button.click()
            print("===================== Парсится сайт 74mv.ru =====================")

        def parsing():
            print("Горный оазис")
            for i in range(4):
                css = f"#bd_results > div.blog > div > div.itemnews.leading-{i} > p:nth-child(2)"
                WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CSS_SELECTOR, css)))
                sale = browser.find_element(By.CSS_SELECTOR, css)
                print(sale.text)

                add_txtstock_in_docx(sale.text)

        get_page()
        parsing()
    # output IMG and text:
    def Chebarkul_source():
        def get_page():
            browser.get("https://chebistok.ru/#sale")
            print("===================== Парсится сайт chebistok.ru =====================")

        def parsing_img():
            print("Чебаркульский исток")

            # IMG 1
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "hidden-600")))
            image1 = browser.find_element(By.CLASS_NAME, "hidden-600")
            url1 = image1.get_attribute("src")
            print("URL1 = ", url1)
            name1 = url1.split('/')[-1]
            print("NAME1 = ", name1)
            urllib.request.urlretrieve(url1, name1)
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "swiper-button-next")))
            next_button = browser.find_element(By.CLASS_NAME, "swiper-button-next")
            next_button.click()
            time.sleep(2)

            # IMG 2
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.CLASS_NAME, "visible-600")))
            image2 = browser.find_element(By.CLASS_NAME, "visible-600")
            url2 = image2.get_attribute("src")
            print("URL2 = ", url2)
            name2 = url2.split('/')[-1]
            print("NAME2 = ", name2)
            urllib.request.urlretrieve(url2, name2)

            # add in .docx file
            add_imgstock_in_docx(name1)
            time.sleep(0.5)
            add_imgstock_in_docx(name2)
            time.sleep(0.5)
        def parsing_text():
            X_Path_title = "/html/body/div[3]/section[7]/div/div/div[3]/div[1]/div/div[2]/p[1]"
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.XPATH, X_Path_title)))
            sale_title = browser.find_element(By.XPATH, X_Path_title)
            time.sleep(1)
            X_Path_price = "/html/body/div[3]/section[7]/div/div/div[3]/div[1]/div/div[2]/p[3]"
            WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.XPATH, X_Path_price)))
            sale_price = browser.find_element(By.XPATH, X_Path_price)
            print(sale_title.text, sale_price.text, sep=" ")
            text_stock = sale_title.text + '\n' + sale_price.text
            add_txtstock_in_docx(text_stock)
            for i in range(1, 6):
                WebDriverWait(browser, 20).until(
                    EC.presence_of_element_located((By.XPATH, "/html/body/div[3]/section[7]/div/div/div[2]")))
                swipe_button = browser.find_element(By.CLASS_NAME, "sldercooler").find_element(By.CLASS_NAME,
                                                                                               "swiper-button-next")
                swipe_button.click()
                time.sleep(3)

                X_Path_title = f"/html/body/div[3]/section[7]/div/div/div[3]/div[{i + 1}]/div/div[2]/p[1]"
                WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.XPATH, X_Path_title)))
                sale_title = browser.find_element(By.XPATH, X_Path_title)
                X_Path_price = f"/html/body/div[3]/section[7]/div/div/div[3]/div[{i + 1}]/div/div[2]/p[3]"
                WebDriverWait(browser, 20).until(EC.presence_of_element_located((By.XPATH, X_Path_price)))
                sale_price = browser.find_element(By.XPATH, X_Path_price)
                print(sale_title.text, sale_price.text, sep=" ")

                text_stock = str(sale_title.text) + " " + str(sale_price.text)
                time.sleep(0.1)
                add_txtstock_in_docx(text_stock)
                time.sleep(2)

        get_page()
        parsing_img()
        parsing_text()


    # добавить акцию в виде текста в Docx файл
    def add_txtstock_in_docx(text_stock):
        file_name = STOCK_FILE_NAME
        doc = docx.Document(file_name)

        # добавляем нужный текст
        par1 = doc.add_paragraph("")
        par1.add_run("Акция\n").bold = True
        par1.add_run(text_stock + '\n')

        doc.save(file_name)
    # добавить изображение акции в Docx файл
    def add_imgstock_in_docx(image_name):
        file_name = STOCK_FILE_NAME
        doc = docx.Document(file_name)

        # добавляем текст АКЦИЯ
        par1 = doc.add_paragraph("")
        par1.add_run("Акция").bold = True
        # добавляем изображение
        doc.add_picture(image_name, width=Cm(10))
        doc.save(file_name)
    # добавить название компании в блок акций в файл "акции.docx"
    def add_company_name(name):
        file_name = STOCK_FILE_NAME
        doc = docx.Document(file_name)
        par1 = doc.add_paragraph("")

        run = par1.add_run(name + "\n")
        run.font.size = Pt(14)
        run.bold = True

        doc.save(file_name)
    # стереть содержимое файла "акции.docx"
    def erase_content():
        file_name = STOCK_FILE_NAME
        doc = docx.Document(file_name)
        para = doc.paragraphs
        for i in para:
            p = i._element
            p.getparent().remove(p)
            i._element = None
        doc.save(file_name)
    # разделить акции по месяцам и тд
    def edit_files_stocks():
        months = ["Январь", "Февраль", "Март", "Апрель", "Май", "Июнь", "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь",
                  "Декабрь"]

        now = datetime.datetime.now()
        date = now.strftime("%d.%m.%Y")
        m_and_y = now.strftime("%m.%Y")  # Текущий месяц и год через точку
        m = now.strftime("%m")  # Текущий месяц
        y = now.strftime("%Y")  # Текущий год

        def edit_month_file():
            # проверяем дату в файле db.txt
            bdfile = open("bd.txt", "r+")
            # если запись есть редактируем файл
            if m_and_y + '\n' in bdfile:
                print("дата найдена в bd.txt")

                # считываем файл "акции _m_ _y_.docx"
                docx_name_old = "акции " + months[int(m) - 1] + " " + y + ".docx"
                print("считываем файл " + docx_name_old)
                doc_m_y_old = docx.Document(docx_name_old)
                all_paras_old = doc_m_y_old.paragraphs
                text_old_doc = ""
                for para in all_paras_old:
                    text_old_doc += para.text

                # считываем файл "акции.docx"(с последним парсингом)
                docx_name_act = "акции.docx"
                doc_actual = docx.Document(docx_name_act)
                all_paras_new = doc_actual.paragraphs

                # проверяем есть ли новые акции
                add_stocks = []  # для дальнейшего добавления в файл "акции _m_ _y_.docx"
                for para in all_paras_new:
                    # если есть, то
                    if para.text not in text_old_doc:
                        print("Следующей записи нет в файле " + docx_name_old + ": \n" + para.text)
                        add_stocks.append(para.text)

                # добавляем новые акции в файл "акции _m_ _y_.docx"
                if len(add_stocks) != 0:
                    for stock in add_stocks:
                        doc_m_y_old.add_paragraph("НОВАЯ ЗАПИСЬ(дата добавления " + date + "):")
                        doc_m_y_old.add_paragraph(stock)
                        doc_m_y_old.save(docx_name_old)
            # добавляем дату в файл db.txt если ее там нет и создаем новый файл
            else:
                print("Дата не найдена в bd.txt")
                print("Добавляем запись даты в bd.txt")
                bdfile.write(m_and_y + '\n')
                # копируем файл "акции.docx" и переименовываем его в "акции _m_ _y_.docx"
                new_name = "акции " + months[int(m) - 1] + " " + y + ".docx"
                shutil.copy("акции.docx", new_name)
                print("создан файл " + new_name)

        # сверяем текующие спрасенные данные с файлом "все акции.docx" и если чего то нет => добавляем их в файл
        def checkfile_allstocks():
            # считываем файл "все акции.docx"
            doc_all_stocks_name = "все акции.docx"
            doc_all_stock = docx.Document(doc_all_stocks_name)
            all_paras_allstocks = doc_all_stock.paragraphs
            all_text_alls = ""
            for para in all_paras_allstocks:
                all_text_alls += para.text

            # считываем файл "акции.docx"
            doc_actual_name = "акции.docx"
            doc_actual = docx.Document(doc_actual_name)
            all_paras_act = doc_actual.paragraphs

            for para in all_paras_act:
                if para.text not in all_text_alls:
                    print("Следующей записи нет в файле \"" + doc_all_stocks_name + "\":\n" + para.text)
                    doc_all_stock.add_paragraph("НОВАЯ ЗАПИСЬ(дата добавления " + date + "):")
                    doc_all_stock.add_paragraph(para.text)
                    doc_all_stock.save(doc_all_stocks_name)

        edit_month_file()
        checkfile_allstocks()

    erase_content()

    browser = webdriver.Chrome()
    browser.maximize_window()

    total_error = 0

    try:
        add_company_name("ЧЕБАРКУЛЬСКИЙ ИСТОЧНИК")
        Chebarkul_source()
    except:
        total_error += 1
        print("-------------Не удалось спарсить chebistok.ru-------------")

    try:
        add_company_name("НИАГАРА")
        niagara()
    except:
        total_error += 1
        print("-------------Не удалось спарсить niagara74.ru-------------")

    try:
        add_company_name("ЛЮКС-ВОДА")
        Suite_water()
    except:
        total_error += 1
        print("-------------Не удалось спарсить l-w.ru-------------")

    try:
        add_company_name("ЖИВАЯ КАПЛЯ")
        parsing_living_capla()
    except:
        total_error += 1
        print("-------------Не удалось спарсить живаякапля.рф-------------")

    try:
        add_company_name("АКВА-МОБИЛЬ")
        parsing_aqua_mobil()
    except:
        total_error += 1
        print("-------------Не удалось спарсить aqua-mobil.ru-------------")

    try:
        add_company_name("АРТЕЗИАНСКАЯ ВОДА")
        artvod()
    except:
        total_error += 1
        print("-------------Не удалось спарсить artvod.ru-------------")

    try:
        add_company_name("ВОДА КРИСТАЛЬНАЯ")
        crystal_water()
    except:
        total_error += 1
        print("-------------Не удалось спарсить voda174.ru-------------")

    try:
        add_company_name("ГОРНЫЙ ОАЗИС")
        mountain_oasis()
    except:
        total_error += 1
        print("-------------Не удалось спарсить www.74mv.ru-------------")

    close_browser()


    # проверяем акции на изменение
    #compare_docx_files()
    #convert("изменения в акциях.docx")

    # создаем копию файла "акции.docx" с именем "акции_old.docx"
    #create_copy_docx()

    if total_error != 0:
        msg = " ================= Всего ошибок парсинга: " + str(total_error) + '\n' + "Попробуйте запустить парсинг еще раз ================= "
        print(msg)
    else:
        print('Акции успешно спарсены, ожидайте...')
        try:
            # Конвертируем(создаём) "акции.docx" в "акции.pdf"
            edit_files_stocks()
            convert("все акции.docx", "все акции.pdf")
        except:
            edit_files_stocks()
            doc = aw.Document("все акции.docx")
            doc.save("все акции.pdf")

#parsing_price()
#parsing_stock()