import shutil
import random
import datetime

from openpyxl import load_workbook  # для записи в excel
from openpyxl.worksheet.dimensions import ColumnDimension, DimensionHolder
from openpyxl.utils import get_column_letter
from openpyxl.styles import PatternFill, Border, Side, Alignment, Font

import docx
from docx.shared import Cm, Pt
from docx.enum.text import WD_COLOR_INDEX

STOCK_FILE_NAME = "акции.docx"
FILE_NAME_PRICE = "цены за всё время.xlsx"


# PRICES, EXCEL
# проверить excel файл на наличие листа, если есть - удалить
def check_sheet_in_excel():
        file_name = FILE_NAME_PRICE
        wbook = load_workbook(file_name)

        now = datetime.datetime.now()
        date = now.strftime("%d.%m.%Y")
        sheet_name = date  # дата с компьютера
        if sheet_name in wbook:
            wbook.remove(wbook[sheet_name])
            wbook.save(filename=file_name)
        wbook.save(filename=file_name)
# функция для добавления данных в Excel
def add_excel(name, price_one_b, price_two_b, price_tara):
        colors = ["FFFFE4B5", "FFFFE4E1", "7FFFD4", "E0EEEE", "FFE4C4", "FFEBCD", "98F5FF", "FF9912", "66CD00",
                  "FF7F24",
                  "FFF8DC", "FF8C00", "BF3EFF", "F2F2F2", "ADFF2F", "FF6A6A", "FFF68F", "FFF0F5", "BFEFFF", "FFA07A",
                  "A4D3EE",
                  "E066FF", "B3EE3A", "BBFFFF", "FFBBFF", "FFC1C1"]

        file_name = FILE_NAME_PRICE
        wbook = load_workbook(file_name)

        # стиль шрифта для названий
        font_name = Font(size=12, bold=True)

        index = len(wbook.sheetnames)

        now = datetime.datetime.now()
        date = now.strftime("%d.%m.%Y")
        sheet_name = date  # дата с компьютера

        def function():
            # ширина колонок
            def len_columns():
                dim_holder = DimensionHolder(worksheet=ws)
                for col in range(ws.min_column, ws.max_column + 1):
                    # width=len(cell_value) * 1.23
                    dim_holder[get_column_letter(col)] = ColumnDimension(ws, min=col, max=col, width=25)
                ws.column_dimensions = dim_holder

            # узнать, какая колонка свободна
            def check_empty_cell():
                empty_cell_randc = []

                value = 1
                row = 1
                column = 1
                while (value != None):
                    value = (ws.cell(row=row, column=column)).value
                    if value == None:
                        break
                    else:
                        column += 1

                empty_cell_randc.append(row)
                empty_cell_randc.append(column)
                return empty_cell_randc

            # добавить значения в лист
            def add_values():
                call_name_r1 = ws.cell(row=call_row, column=call_col).column_letter + str(
                    ws.cell(row=call_row, column=call_col).row)
                ws.cell(row=call_row, column=call_col, value=name).font = font_name
                ws.cell(row=call_row, column=call_col, value=name).border = border_h
                ws.cell(row=call_row, column=call_col).alignment = Alignment(wrap_text=True, horizontal='center',
                                                                             vertical='center')  # включаем перенос строк и выравнивание для ячеек в первой строке

                # ws.cell(row=call_row + 1, column=call_col, value=price_one_b).number_format = '#,##0 ₽'
                ws.cell(row=call_row + 1, column=call_col, value=price_one_b).border = border_2

                ws.cell(row=call_row + 2, column=call_col, value=price_two_b).border = border_2

                # формула  пример: =ЕСЛИ(B3 = 0; 0;1-'02.18'!B3/B3)     АНГЛОЯЗЫЧНЫЙ ВАРИАНТ, там не ; а ,
                # =ЕСЛИ('31.10.2022'!B1='06.11.2022'!B1;ЕСЛИ(B3=0; 0; 1-'30.10.2022'!B3/B3);0)
                # название ячейки строка 3
                call_name_r3 = ws.cell(row=call_row + 2, column=call_col).column_letter + str(
                    ws.cell(row=call_row + 2, column=call_col).row)
                # print(call_name_r3)
                # print("=IF(" + '\'' + name_last_sheet + "\'!" + call_name_r1 + "=\'" + sheet_name + "\'!" + call_name_r1 + ', IF(' + call_name_r3 + "=0, 0, 1-"+"\'" + name_last_sheet + "\'" + "!" + call_name_r3 + "/" + call_name_r3 + ")" + ", 0)")

                ws.cell(row=call_row + 3, column=call_col,
                        value="=IF(" + '\'' + name_last_sheet + "\'!" + call_name_r1 + "=\'" + sheet_name + "\'!" + call_name_r1 + ', IF(' + call_name_r3 + "=0, 0, 1-" + "\'" + name_last_sheet + "\'" + "!" + call_name_r3 + "/" + call_name_r3 + ")" + ", 0)")
                ws.cell(row=call_row + 3, column=call_col).border = border_2
                ws.cell(row=call_row + 3, column=call_col).number_format = "0.0%"

                ws.cell(row=call_row + 4, column=call_col, value=price_tara).border = border_3

            # настроить высоту первой строки
            ws.row_dimensions[1].height = 64

            call_row = check_empty_cell()[0]
            call_col = check_empty_cell()[1]

            # вписываем текст в самые левые ячейки
            if call_row == 1 and call_col == 1:
                # назначаем цвет первой колонке
                h_fill = PatternFill(start_color='FFE4C4',
                                     end_color='FFE4C4',
                                     fill_type='solid')
                ws['A1'].fill = h_fill
                ws['A2'].fill = h_fill
                ws['A3'].fill = h_fill
                ws['A4'].fill = h_fill
                ws['A5'].fill = h_fill

                # вписываем текст в первые пять ячеек первой колонки
                ws.cell(row=call_row, column=call_col, value="Название").font = font_name
                ws.cell(row=call_row, column=call_col).border = border_h
                ws.cell(row=call_row, column=call_col).alignment = Alignment(wrap_text=True, horizontal='center',
                                                                             vertical='center')  # включаем перенос строк и выравнивание для ячеек в первой строке

                ws.cell(row=call_row + 1, column=call_col, value="19л 1шт").border = border_h
                ws.cell(row=call_row + 1, column=call_col).font = font_name

                ws.cell(row=call_row + 2, column=call_col, value="19л от 2х").border = border_h
                ws.cell(row=call_row + 2, column=call_col).font = font_name

                ws.cell(row=call_row + 3, column=call_col, value="Изм. цены").border = border_h
                ws.cell(row=call_row + 3, column=call_col).font = font_name

                ws.cell(row=call_row + 4, column=call_col, value="Тара").border = border_h
                ws.cell(row=call_row + 4, column=call_col).font = font_name

                call_col += 1

            add_values()

            len_columns()
            ws.column_dimensions['A'].width = 15
            ws.alignment = Alignment(wrap_text=True)
            wbook.save(file_name)

        def color_sheet():
            # цвет листа внизу
            sheet_properties = ws.sheet_properties
            random_num = random.randint(0, len(colors) - 1)
            sheet_properties.tabColor = colors[random_num]

        # настроить стили граней ячеек
        bd = Side(border_style='medium')
        bd2 = Side(border_style='thin')
        border_h = Border(left=bd, top=bd, right=bd, bottom=bd)
        border_2 = Border(left=bd, top=bd2, right=bd, bottom=bd2)
        border_3 = Border(left=bd, top=bd2, right=bd, bottom=bd)

        if sheet_name in wbook:
            print("Такая страница есть")
            ws = wbook[sheet_name]  # лист
            # узнаём имя листа с предыдущими измерениями
            name_last_sheet = wbook.sheetnames[index - 2]
            color_sheet()
            function()

        else:
            print("Такой страницы нет")
            wbook.create_sheet(index=index, title=sheet_name)
            ws = wbook[sheet_name]  # лист
            name_last_sheet = wbook.sheetnames[index - 1]
            color_sheet()
            function()
# добавить цену на товар в сводный лист в Excel
def add_listsummary(tovar_name, price):
        file_name = FILE_NAME_PRICE
        sheet_name = "Общая от 2-х"

        wbook = load_workbook(file_name)
        sheet = wbook[sheet_name]
        print("OK, лист найден")

        now = datetime.datetime.now()
        date = now.strftime("%d.%m.%Y")
        print("NOW DATE = " + date)

        # проверка есть ли колонка с названием товара:  return(True / False)
        def check_item_name():
            check_item_name = True
            print("Ищем колонку с названием товара...")

            value = "name"
            row = 1
            column = 2
            while (value != tovar_name and value != None):
                value = (sheet.cell(row=row, column=column)).value
                if value == tovar_name:
                    return check_item_name
                else:
                    column += 1

            check_item_name = False
            return check_item_name

        # ПОИСК НУЖНОЙ КОЛОНКИ
        def find_desired_column():
            print("Ищем нужную колонку")

            value = "name"
            row = 1
            column = 2
            while (value != tovar_name and value != None):
                value = (sheet.cell(row=row, column=column)).value
                if value == tovar_name:
                    print("Нужная колонка найдена, её имя: ", value, '\n')
                    return column
                else:
                    column += 1

            return column

        # ПОИСК НУЖНОЙ СТРОЧКИ В СТОЛБЦЕ
        def find_desired_row():
            print("Ищем нужную строчку...")
            row = 2
            column = 1
            value = (sheet.cell(row=row, column=column)).value

            while (value != date and value != None):
                value = (sheet.cell(row=row, column=column)).value
                #print(value, '\n')
                if (str(value) == date):
                    return row
                # если строчка с сегодняшней датой не найдена - записать ее в первый столбец
                if (value == None):
                    sheet.cell(row=row, column=1, value=date)
                    return row
                row += 1

            sheet.cell(row=row, column=1, value=date)
            return row

        item_column = find_desired_column()
        empty_row = find_desired_row()

        # если колонка с таким именем товара уже есть
        if check_item_name():
            print("Такой товар есть : )")
            # Добавляем значение в ячейку нужного столбца
            sheet.cell(row=empty_row, column=item_column, value=price)
        # если такой колонки нет
        else:
            print("Такого товара не было")
            # Добавляем значение в ячейку нужного столбца
            sheet.cell(row=1, column=item_column - 1, value=tovar_name)
            # добавляем название товара в первую строчку таблицы
            sheet.cell(row=empty_row, column=item_column - 1, value=price)

        wbook.save(file_name)


# STOCK, WORD, PDF
# добавить акцию в виде текста в Docx файл
def add_txtstock_in_docx(text_stock):
        file_name = STOCK_FILE_NAME
        doc = docx.Document(file_name)

        # добавляем нужный текст
        par1 = doc.add_paragraph("")
        par1.add_run("Акция\n").bold = True
        par1.add_run(text_stock + '\n')

        doc.save(file_name)
# добавить изображение акции в Docx файл
def add_imgstock_in_docx(image_name):
        file_name = STOCK_FILE_NAME
        doc = docx.Document(file_name)

        # добавляем текст АКЦИЯ
        par1 = doc.add_paragraph("")
        par1.add_run("Акция").bold = True
        # добавляем изображение
        doc.add_picture(image_name, width=Cm(10))
        doc.save(file_name)
# добавить название компании в блок акций в файл "акции.docx"
def add_company_name(name):
        file_name = STOCK_FILE_NAME
        doc = docx.Document(file_name)
        par1 = doc.add_paragraph("")

        run = par1.add_run(name + "\n")
        run.font.size = Pt(14)
        run.bold = True

        doc.save(file_name)
# стереть содержимое файла "акции.docx"
def erase_content():
        file_name = STOCK_FILE_NAME
        doc = docx.Document(file_name)
        para = doc.paragraphs
        for i in para:
            p = i._element
            p.getparent().remove(p)
            i._element = None
        doc.save(file_name)
# разделить акции по месяцам и тд
def edit_files_stocks():
        print("Собираем файл...")
        months = ["Январь", "Февраль", "Март", "Апрель", "Май", "Июнь", "Июль", "Август", "Сентябрь", "Октябрь", "Ноябрь",
                  "Декабрь"]

        now = datetime.datetime.now()
        date = now.strftime("%d.%m.%Y")
        m_and_y = now.strftime("%m.%Y")  # Текущий месяц и год через точку
        m = now.strftime("%m")  # Текущий месяц
        y = now.strftime("%Y")  # Текущий год

        def edit_month_file():
            # проверяем дату в файле db.txt
            bdfile = open("bd.txt", "r+")
            # если запись есть редактируем файл
            if m_and_y + '\n' in bdfile:
                print("дата найдена в bd.txt")

                # считываем файл "акции _m_ _y_.docx"
                docx_name_old = "акции " + months[int(m) - 1] + " " + y + ".docx"
                print("считываем файл " + docx_name_old)
                doc_m_y_old = docx.Document(docx_name_old)
                all_paras_old = doc_m_y_old.paragraphs
                text_old_doc = ""
                for para in all_paras_old:
                    text_old_doc += para.text

                # считываем файл "акции.docx"(с последним парсингом)
                docx_name_act = "акции.docx"
                doc_actual = docx.Document(docx_name_act)
                all_paras_new = doc_actual.paragraphs

                # проверяем есть ли новые акции
                add_stocks = []  # для дальнейшего добавления в файл "акции _m_ _y_.docx"
                for para in all_paras_new:
                    # если есть, то
                    if para.text not in text_old_doc:
                        print("Следующей записи нет в файле " + docx_name_old + ": \n" + para.text)
                        add_stocks.append(para.text)

                # добавляем новые акции в файл "акции _m_ _y_.docx"
                if len(add_stocks) != 0:
                    for stock in add_stocks:
                        doc_m_y_old.add_paragraph("НОВАЯ ЗАПИСЬ(дата добавления " + date + "):")
                        doc_m_y_old.add_paragraph(stock)
                        doc_m_y_old.save(docx_name_old)
            # добавляем дату в файл db.txt если ее там нет и создаем новый файл
            else:
                print("Дата не найдена в bd.txt")
                print("Добавляем запись даты в bd.txt")
                bdfile.write(m_and_y + '\n')
                # копируем файл "акции.docx" и переименовываем его в "акции _m_ _y_.docx"
                new_name = "акции " + months[int(m) - 1] + " " + y + ".docx"
                shutil.copy("акции.docx", new_name)
                print("создан файл " + new_name)

        # сверяем текующие спрасенные данные с файлом "все акции.docx" и если чего то нет => добавляем их в файл

        def checkfile_allstocks():
            # считываем файл "все акции.docx"
            doc_all_stocks_name = "все акции.docx"
            doc_all_stock = docx.Document(doc_all_stocks_name)
            all_paras_allstocks = doc_all_stock.paragraphs
            all_text_alls = ""
            for para in all_paras_allstocks:
                all_text_alls += para.text

            # считываем файл "акции.docx"
            doc_actual_name = "акции.docx"
            doc_actual = docx.Document(doc_actual_name)
            all_paras_act = doc_actual.paragraphs

            bnewstock = 0
            for para in all_paras_act:
                if para.text not in all_text_alls:
                    bnewstock = 1
                    print("Следующей записи нет в файле \"" + doc_all_stocks_name + "\":\n" + para.text)
                    # также выделяем маркером
                    p = doc_all_stock.add_paragraph()
                    p.add_run("НОВАЯ ЗАПИСЬ(дата добавления " + date + "):").font.highlight_color = WD_COLOR_INDEX.YELLOW
                    doc_all_stock.add_paragraph(para.text)
                    doc_all_stock.save(doc_all_stocks_name)
            if bnewstock == 1:
                return 1
            else:
                return 0

        edit_month_file()
        if checkfile_allstocks() == 1:
            return 1
        else:
            return 0